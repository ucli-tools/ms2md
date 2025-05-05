"""
DOCX processing module for MS2MD.

This module provides functions for extracting and processing content from
Microsoft Word (.docx) documents.
"""

import os
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple, Union

import docx
from docx.document import Document
# The math module is not available in the current version of python-docx
# from docx.oxml.math import CT_OMath
from docx.oxml.text.paragraph import CT_P
from docx.text.paragraph import Paragraph

from ms2md.utils.logging_utils import get_logger

logger = get_logger(__name__)


def extract_docx_content(
    docx_path: Union[str, Path]
) -> Dict[str, Any]:
    """
    Extract content from a Word document.
    
    Args:
        docx_path: Path to the Word document
        
    Returns:
        Dictionary containing extracted content
    """
    doc_path = Path(docx_path)
    
    if not doc_path.exists():
        raise FileNotFoundError(f"Document not found: {doc_path}")
    
    logger.info(f"Extracting content from {doc_path}")
    
    try:
        doc = docx.Document(doc_path)
    except Exception as e:
        logger.error(f"Failed to open document {doc_path}: {str(e)}")
        raise ValueError(f"Failed to open document: {str(e)}")
    
    # Extract document properties
    properties = extract_document_properties(doc)
    
    # Extract paragraphs
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append({
                "text": para.text,
                "style": para.style.name if para.style else "Normal",
            })
    
    # Extract tables
    tables = []
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text)
            table_data.append(row_data)
        tables.append(table_data)
    
    # Extract equations
    equations = extract_equations(doc)
    
    # Extract images
    images = extract_image_info(doc_path)
    
    return {
        "properties": properties,
        "paragraphs": paragraphs,
        "tables": tables,
        "equations": equations,
        "images": images,
    }


def extract_document_properties(doc: Document) -> Dict[str, str]:
    """
    Extract document properties from a Word document.
    
    Args:
        doc: Word document object
        
    Returns:
        Dictionary of document properties
    """
    properties = {}
    
    # Core properties
    if hasattr(doc, "core_properties"):
        core_props = doc.core_properties
        if hasattr(core_props, "title") and core_props.title:
            properties["title"] = core_props.title
        if hasattr(core_props, "author") and core_props.author:
            properties["author"] = core_props.author
        if hasattr(core_props, "created") and core_props.created:
            properties["created"] = str(core_props.created)
        if hasattr(core_props, "modified") and core_props.modified:
            properties["modified"] = str(core_props.modified)
        if hasattr(core_props, "subject") and core_props.subject:
            properties["subject"] = core_props.subject
        if hasattr(core_props, "keywords") and core_props.keywords:
            properties["keywords"] = core_props.keywords
    
    return properties


def extract_equations(doc: Document) -> List[Dict[str, Any]]:
    """
    Extract equations from a Word document.
    
    Args:
        doc: Word document object
        
    Returns:
        List of dictionaries containing equation information
    """
    equations = []
    
    # Iterate through paragraphs to find equations
    for i, para in enumerate(doc.paragraphs):
        # In newer versions of python-docx, we can't use namespaces in xpath
        # Try to find equations by looking for specific patterns in the XML
        try:
            # Try to find math elements without using namespaces
            math_elements = []
            
            # Check if paragraph contains any math-related content
            if "oMath" in para._element.xml:
                # Just add a placeholder for now
                equations.append({
                    "type": "omml",
                    "paragraph_index": i,
                    "xml": "<placeholder>",
                    "display": False,  # Default to inline equation
                })
        except Exception as e:
            logger.warning(f"Error extracting equation from paragraph {i}: {str(e)}")
    
    logger.info(f"Extracted {len(equations)} equations")
    return equations


def is_display_equation(para: Paragraph, omml) -> bool:
    """
    Determine if an equation is a display equation (centered on its own line).
    
    Args:
        para: Paragraph containing the equation
        omml: OMML equation element
        
    Returns:
        True if it's a display equation, False otherwise
    """
    # Check if there's significant text before or after the equation
    text_content = para.text.strip()
    if not text_content or text_content.isspace():
        return True
    
    return False


def extract_image_info(docx_path: Union[str, Path]) -> List[Dict[str, Any]]:
    """
    Extract information about images in a Word document.
    
    Args:
        docx_path: Path to the Word document
        
    Returns:
        List of dictionaries containing image information
    """
    # This is a simplified version - extracting actual images requires
    # more complex handling of the DOCX ZIP structure
    doc = docx.Document(docx_path)
    
    images = []
    image_parts = []
    
    # Get all image parts from the document
    for rel in doc.part.rels.values():
        if rel.reltype == 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/image':
            image_parts.append(rel.target_part)
    
    # Create image info dictionaries
    for i, img_part in enumerate(image_parts):
        # Get content type and extension
        content_type = img_part.content_type
        ext = content_type.split('/')[-1]
        if ext == 'jpeg':
            ext = 'jpg'
        
        images.append({
            "index": i,
            "content_type": content_type,
            "extension": ext,
            "size_bytes": len(img_part.blob),
        })
    
    logger.info(f"Found {len(images)} images")
    return images